import tkinter as tk
from tkinter import ttk, scrolledtext, messagebox, filedialog
import threading
import json
import os
from datetime import datetime
from Chatbot import ChatBot
import tempfile

# Importar librerías para leer diferentes tipos de archivos
try:
    from docx import Document
    DOCX_DISPONIBLE = True
except ImportError:
    DOCX_DISPONIBLE = False
    Document = None

try:
    import PyPDF2
    PDF_DISPONIBLE = True
except ImportError:
    PDF_DISPONIBLE = False
    PyPDF2 = None

# Importar Pillow para manejar imágenes del portapapeles
try:
    from PIL import Image, ImageGrab
    PIL_DISPONIBLE = True
except ImportError:
    PIL_DISPONIBLE = False
    Image = None
    ImageGrab = None

class ChatBotGUI:
    def __init__(self):
        self.ventana = tk.Tk()
        self.ventana.title("🤖 Asistente Virtual AI - Análisis Inteligente de Documentos")
        self.ventana.geometry("900x700")
        self.ventana.minsize(800, 600)
        
        # Configurar icono de la ventana (si tienes un archivo .ico)
        try:
            # self.ventana.iconbitmap("icono.ico")  # Descomenta si tienes icono
            pass
        except:
            pass
        
        # Configurar modo tema (True = oscuro, False = claro)
        self.modo_oscuro = False
        
        self.ventana.configure(bg='#f8f9fa')
        
        # Configurar el icono y estilo
        self.configurar_estilo()
        
        # Inicializar el chatbot
        self.chatbot = ChatBot("Asistente Virtual")

        # Crear la interfaz
        self.crear_interfaz()
        
        # Mensaje de bienvenida mejorado
        self.mostrar_mensaje_bot("🎉 ¡Hola! Soy tu Asistente Virtual con IA avanzada. ¿En qué puedo ayudarte hoy? \n\n✨ Puedo analizar documentos, generar casos de prueba, crear manuales de usuario y mucho más. ¡Adjunta archivos o simplemente pregúntame!")

        # Configurar evento de cierre
        self.ventana.protocol("WM_DELETE_WINDOW", self.cerrar_aplicacion)
    
    def configurar_estilo(self):
        """Configura el estilo visual de la aplicación"""
        self.style = ttk.Style()
        self.style.theme_use('clam')
        
        # Definir temas con colores más atractivos y gradientes
        self.temas = {
            'claro': {
                'fondo_principal': '#ffffff',
                'fondo_ventana': '#ffffff',
                'fondo_chat': '#ffffff',  # Fondo blanco limpio como ChatGPT
                'fondo_entry': '#f9f9f9',
                'fondo_frame': '#ffffff',
                'texto_principal': '#374151',
                'texto_secundario': '#6b7280',
                'usuario': '#f7f7f8',  # Fondo sutil para usuario
                'bot': '#ffffff',      # Fondo blanco para bot
                'texto_usuario': '#374151',
                'texto_bot': '#374151',
                'borde': '#e5e7eb',
                'boton_bg': '#3b82f6',
                'boton_fg': '#ffffff',
                'boton_hover': '#2563eb',
                'boton_active': '#1d4ed8',
                'accent': '#ef4444',
                'success': '#10b981',
                'warning': '#f59e0b',
                'tree_bg': '#ffffff',
                'tree_fg': '#374151',
                'tree_select': '#3b82f6'
            },
            'oscuro': {
                'fondo_principal': '#212121',
                'fondo_ventana': '#212121',
                'fondo_chat': '#343541',   # Fondo como ChatGPT modo oscuro
                'fondo_entry': '#40414f',
                'fondo_frame': '#212121',
                'texto_principal': '#ffffff',
                'texto_secundario': '#9ca3af',
                'usuario': '#2f2f2f',      # Fondo sutil para usuario
                'bot': '#444654',          # Fondo diferente para bot
                'texto_usuario': '#ffffff',
                'texto_bot': '#ffffff',
                'borde': '#4b5563',
                'boton_bg': '#6366f1',
                'boton_fg': '#ffffff',
                'boton_hover': '#4f46e5',
                'boton_active': '#4338ca',
                'accent': '#f87171',
                'success': '#34d399',
                'warning': '#fbbf24',
                'tree_bg': '#444654',
                'tree_fg': '#ffffff',
                'tree_select': '#6366f1'
            }
        }
        
        self.aplicar_tema()
        self.configurar_estilos_dinamicos()
    
    def aplicar_tema(self):
        """Aplica el tema actual a todos los componentes"""
        tema_actual = 'oscuro' if self.modo_oscuro else 'claro'
        colores = self.temas[tema_actual]
        
        # Configurar ventana principal
        self.ventana.configure(bg=colores['fondo_principal'])
        
        # Configurar estilos TTK
        self.style.configure('TFrame', 
                           background=colores['fondo_frame'],
                           bordercolor=colores['borde'])
        
        self.style.configure('TLabel', 
                           background=colores['fondo_frame'], 
                           foreground=colores['texto_principal'],
                           font=('Segoe UI', 9))
        
        self.style.configure('TButton', 
                           background=colores['boton_bg'], 
                           foreground=colores['boton_fg'],
                           borderwidth=1,
                           focuscolor='none',
                           font=('Segoe UI', 9))
        
        self.style.map('TButton',
                      background=[('active', colores['borde']),
                                ('pressed', colores['texto_secundario'])])
        
        self.style.configure('TLabelFrame', 
                           background=colores['fondo_frame'], 
                           foreground=colores['texto_principal'],
                           bordercolor=colores['boton_bg'],
                           borderwidth=2,
                           relief='solid')
        
        self.style.configure('TLabelFrame.Label', 
                           background=colores['fondo_frame'], 
                           foreground=colores['boton_bg'],
                           font=('Segoe UI', 10, 'bold'))
        
        self.style.configure('TEntry',
                           fieldbackground=colores['fondo_entry'],
                           foreground=colores['texto_principal'],
                           bordercolor=colores['borde'],
                           lightcolor=colores['borde'],
                           darkcolor=colores['borde'],
                           font=('Segoe UI', 10))
        
        # Configurar Treeview
        self.style.configure('Treeview',
                           background=colores['tree_bg'],
                           foreground=colores['tree_fg'],
                           fieldbackground=colores['tree_bg'],
                           borderwidth=0,
                           font=('Segoe UI', 9))
        
        self.style.configure('Treeview.Heading',
                           background=colores['boton_bg'],
                           foreground=colores['texto_principal'],
                           font=('Segoe UI', 9, 'bold'))
        
        self.style.map('Treeview',
                      background=[('selected', colores['tree_select'])])
        
        # Guardar colores actuales para componentes personalizados
        self.colores = colores
    
    def configurar_estilos_dinamicos(self):
        """Configura estilos dinámicos y animaciones"""
        # Estilo para botones principales con efectos
        self.style.configure('Dynamic.TButton',
                           background=self.colores['boton_bg'],
                           foreground=self.colores['boton_fg'],
                           borderwidth=0,
                           focuscolor='none',
                           font=('Segoe UI', 10, 'bold'),
                           padding=(15, 10))
        
        # Estilo para botón de enviar (más destacado)
        self.style.configure('Send.TButton',
                           background=self.colores['success'],
                           foreground='white',
                           borderwidth=0,
                           focuscolor='none',
                           font=('Segoe UI', 11, 'bold'),
                           padding=(20, 12))
        
        # Estilo para botón de tema
        self.style.configure('Theme.TButton',
                           background=self.colores['accent'],
                           foreground='white',
                           borderwidth=0,
                           focuscolor='none',
                           font=('Segoe UI', 9, 'bold'),
                           padding=(12, 8))
        
        # Estilo para botones de acción
        self.style.configure('Action.TButton',
                           background=self.colores['boton_bg'],
                           foreground=self.colores['boton_fg'],
                           borderwidth=0,
                           focuscolor='none',
                           font=('Segoe UI', 9),
                           padding=(10, 6))
        
        # Configurar mapas de estado para efectos hover
        self.style.map('Dynamic.TButton',
                      background=[('active', self.colores['boton_hover']),
                                ('pressed', self.colores['boton_active'])])
        
        self.style.map('Send.TButton',
                      background=[('active', '#229954'),
                                ('pressed', '#1e8449')])
        
        self.style.map('Theme.TButton',
                      background=[('active', '#c0392b'),
                                ('pressed', '#a93226')])
        
        self.style.map('Action.TButton',
                      background=[('active', self.colores['boton_hover']),
                                ('pressed', self.colores['boton_active'])])
    
    def crear_interfaz(self):
        """Crea todos los elementos de la interfaz"""
        # Frame principal
        frame_principal = ttk.Frame(self.ventana, padding="10")
        frame_principal.grid(row=0, column=0, sticky=(tk.W, tk.E, tk.N, tk.S))
        
        # Configurar peso de filas y columnas
        self.ventana.columnconfigure(0, weight=1)
        self.ventana.rowconfigure(0, weight=1)
        frame_principal.columnconfigure(0, weight=1)
        frame_principal.rowconfigure(1, weight=1)
        
        # Título con estilo moderno
        frame_titulo = ttk.Frame(frame_principal)
        frame_titulo.grid(row=0, column=0, pady=(0, 20), sticky=(tk.W, tk.E))
        frame_titulo.columnconfigure(1, weight=1)
        
        # Icono principal animado
        titulo_icono = ttk.Label(
            frame_titulo, 
            text="🤖", 
            font=('Segoe UI Emoji', 24)
        )
        titulo_icono.grid(row=0, column=0, padx=(0, 10))
        
        # Título principal
        titulo = ttk.Label(
            frame_titulo, 
            text="Asistente Virtual AI", 
            font=('Segoe UI', 22, 'bold'),
            foreground=self.colores['boton_bg']
        )
        titulo.grid(row=0, column=1, sticky=(tk.W))
        
        # Subtítulo
        subtitulo = ttk.Label(
            frame_titulo, 
            text="🚀 Análisis inteligente de documentos con IA avanzada", 
            font=('Segoe UI', 11),
            foreground=self.colores['texto_secundario']
        )
        subtitulo.grid(row=1, column=1, sticky=(tk.W), pady=(5, 0))
        
        # Indicador de estado IA en el título
        self.indicador_ia = ttk.Label(
            frame_titulo,
            text="✨ IA Activa" if self.chatbot.usar_ia else "⚡ Modo Local",
            font=('Segoe UI', 9, 'bold'),
            foreground=self.colores['success'] if self.chatbot.usar_ia else self.colores['warning']
        )
        self.indicador_ia.grid(row=0, column=2, padx=(10, 0), sticky=(tk.E))
        
        # Área de chat
        self.crear_area_chat(frame_principal)
        
        # Área de entrada
        self.crear_area_entrada(frame_principal)
        
        # Área de archivos adjuntos
        self.crear_area_adjuntos(frame_principal)
        
        # Botones adicionales
        self.crear_botones(frame_principal)
        
        # Status bar
        self.crear_status_bar(frame_principal)
    
    def crear_area_chat(self, parent):
        """Crea el área donde se muestran los mensajes"""
        # Frame para el chat
        frame_chat = ttk.LabelFrame(parent, text="💬 Conversación", padding="5")
        frame_chat.grid(row=1, column=0, sticky=(tk.W, tk.E, tk.N, tk.S), pady=(0, 10))
        frame_chat.columnconfigure(0, weight=1)
        frame_chat.rowconfigure(0, weight=1)
        
        # Área de texto con scroll
        self.area_chat = scrolledtext.ScrolledText(
            frame_chat,
            wrap=tk.WORD,
            state=tk.DISABLED,
            font=('Segoe UI', 10),
            bg=self.colores['fondo_chat'],
            fg=self.colores['texto_principal'],
            relief=tk.FLAT,
            borderwidth=1,
            insertbackground=self.colores['texto_principal'],
            selectbackground=self.colores['borde']
        )
        self.area_chat.grid(row=0, column=0, sticky=(tk.W, tk.E, tk.N, tk.S))
        
        self.configurar_tags_chat()
    
    def configurar_tags_chat(self):
        """Configura los tags de colores para el chat"""
        # Configurar tags para mensajes del usuario (estilo ChatGPT)
        if self.modo_oscuro:
            # Modo oscuro: fondo gris para usuario
            fondo_usuario = '#2d2d2d'
            texto_usuario = '#ffffff'
        else:
            # Modo claro: fondo gris claro para usuario
            fondo_usuario = '#f7f7f8'
            texto_usuario = '#343541'
        
        self.area_chat.tag_configure(
            "usuario",
            foreground=texto_usuario,
            background=fondo_usuario,
            font=('Segoe UI', 10, 'bold'),  # Negrita para el usuario
            lmargin1=20,
            lmargin2=20,
            rmargin=20,
            spacing1=8,
            spacing3=8,
            justify='left'
        )
        
        # Configurar tags para mensajes del bot (SIN fondo adicional)
        if self.modo_oscuro:
            texto_bot = '#ffffff'
        else:
            texto_bot = '#374151'
        
        self.area_chat.tag_configure(
            "bot",
            foreground=texto_bot,
            # NO background - usa el fondo del área de chat
            font=('Segoe UI', 10),  # Normal, sin negrita
            lmargin1=20,
            lmargin2=20,
            rmargin=20,
            spacing1=8,
            spacing3=8,
            justify='left'
        )
        
        self.area_chat.tag_configure(
            "timestamp", 
            foreground=self.colores['texto_secundario'], 
            font=('Segoe UI', 8),
            justify='left'
        )
    
    def crear_area_entrada(self, parent):
        """Crea el área donde el usuario escribe mensajes"""
        frame_entrada = ttk.Frame(parent)
        frame_entrada.grid(row=2, column=0, sticky=(tk.W, tk.E), pady=(0, 10))
        frame_entrada.columnconfigure(0, weight=1)
        frame_entrada.rowconfigure(0, weight=1)
        
        # Frame para el área de texto de entrada
        frame_texto = ttk.Frame(frame_entrada)
        frame_texto.grid(row=0, column=0, sticky=(tk.W, tk.E, tk.N, tk.S), padx=(0, 10))
        frame_texto.columnconfigure(0, weight=1)
        frame_texto.rowconfigure(0, weight=1)
        
        # Text widget para escribir mensajes (permite múltiples líneas)
        self.entrada_texto = tk.Text(
            frame_texto,
            font=('Segoe UI', 11),
            height=3,  # Altura inicial de 3 líneas
            width=50,
            wrap=tk.WORD,
            bg=self.colores['fondo_entry'],
            fg=self.colores['texto_principal'],
            insertbackground=self.colores['texto_principal'],
            borderwidth=1,
            relief=tk.SOLID,
            undo=True,  # Habilitar funcionalidad de deshacer
            maxundo=50  # Máximo de 50 operaciones de deshacer
        )
        self.entrada_texto.grid(row=0, column=0, sticky=(tk.W, tk.E), padx=(0, 5))
        
        # Scrollbar para el área de entrada
        scroll_entrada = ttk.Scrollbar(frame_texto, orient=tk.VERTICAL, command=self.entrada_texto.yview)
        self.entrada_texto.configure(yscrollcommand=scroll_entrada.set)
        scroll_entrada.grid(row=0, column=1, sticky=(tk.N, tk.S))
        
        # Botón adjuntar archivos con estilo
        self.boton_adjuntar = ttk.Button(
            frame_entrada,
            text="📎 Adjuntar",
            command=self.adjuntar_archivo,
            style='Action.TButton'
        )
        self.boton_adjuntar.grid(row=0, column=1, padx=(0, 10), sticky=(tk.N))
        
        # Botón enviar con estilo destacado
        self.boton_enviar = ttk.Button(
            frame_entrada,
            text="� Enviar",
            command=self.enviar_mensaje,
            style='Send.TButton'
        )
        self.boton_enviar.grid(row=0, column=2, sticky=(tk.N))
        
        # Bind teclas
        self.entrada_texto.bind('<Return>', self.manejar_enter)
        self.entrada_texto.bind('<Shift-Return>', self.manejar_shift_enter)
        
        # Atajos de teclado estándar
        self.configurar_atajos_teclado()
        
        # Menú contextual (clic derecho)
        self.configurar_menu_contextual()
        
        self.entrada_texto.focus()
        
        # Etiqueta de ayuda con mejor estilo
        label_ayuda = ttk.Label(
            frame_entrada,
            text="💡 Enter: Enviar | Shift+Enter: Nueva línea | Ctrl+V: Pegar imagen | Ctrl+Shift+L: Limpiar",
            font=('Segoe UI', 9, 'italic'),
            foreground=self.colores['texto_secundario']
        )
        label_ayuda.grid(row=1, column=0, columnspan=3, sticky=(tk.W), pady=(8, 0))
    
    def crear_area_adjuntos(self, parent):
        """Crea el área para mostrar archivos adjuntos"""
        self.frame_adjuntos = ttk.LabelFrame(parent, text="📎 Archivos Adjuntos", padding="5")
        self.frame_adjuntos.grid(row=3, column=0, sticky=(tk.W, tk.E), pady=(0, 10))
        self.frame_adjuntos.columnconfigure(0, weight=1)
        
        # Lista para mostrar archivos adjuntos
        self.lista_adjuntos = tk.Listbox(
            self.frame_adjuntos,
            height=3,
            font=('Segoe UI', 9),
            bg=self.colores['fondo_chat'],
            fg=self.colores['texto_principal'],
            selectbackground=self.colores['borde']
        )
        self.lista_adjuntos.grid(row=0, column=0, sticky=(tk.W, tk.E), padx=(0, 10))
        
        # Frame para botón de adjuntos
        frame_boton_adj = ttk.Frame(self.frame_adjuntos)
        frame_boton_adj.grid(row=0, column=1, sticky=(tk.N))
        
        # Botón limpiar todos (único botón)
        self.boton_limpiar_adj = ttk.Button(
            frame_boton_adj,
            text="🧹 Limpiar Adjuntos",
            command=self.limpiar_adjuntos_y_texto
        )
        self.boton_limpiar_adj.pack()
        
        # Lista para almacenar rutas de archivos
        self.archivos_adjuntos = []
        
        # Ocultar el frame inicialmente
        self.frame_adjuntos.grid_remove()
    
    def crear_botones(self, parent):
        """Crea botones adicionales con estilo moderno"""
        frame_botones = ttk.Frame(parent)
        frame_botones.grid(row=4, column=0, sticky=(tk.W, tk.E), pady=(15, 10))
        
        # Crear dos filas de botones para mejor organización
        frame_botones_top = ttk.Frame(frame_botones)
        frame_botones_top.pack(fill=tk.X, pady=(0, 10))
        
        frame_botones_bottom = ttk.Frame(frame_botones)
        frame_botones_bottom.pack(fill=tk.X)
        
        # Primera fila - Botones principales
        boton_limpiar = ttk.Button(
            frame_botones_top,
            text="🧹 Limpiar Chat",
            command=self.limpiar_chat,
            style='Dynamic.TButton'
        )
        boton_limpiar.pack(side=tk.LEFT, padx=(0, 15))
        
        boton_historial = ttk.Button(
            frame_botones_top,
            text="� Historial",
            command=self.abrir_ventana_historial,
            style='Dynamic.TButton'
        )
        boton_historial.pack(side=tk.LEFT, padx=(0, 15))
        
        boton_guardar = ttk.Button(
            frame_botones_top,
            text="💾 Guardar",
            command=self.guardar_conversacion_actual,
            style='Dynamic.TButton'
        )
        boton_guardar.pack(side=tk.LEFT, padx=(0, 15))
        
        # Botón cambiar tema con animación
        self.boton_tema = ttk.Button(
            frame_botones_top,
            text="🌙 Modo Oscuro" if not self.modo_oscuro else "☀️ Modo Claro",
            command=self.cambiar_tema_animado,
            style='Theme.TButton'
        )
        self.boton_tema.pack(side=tk.RIGHT, padx=(15, 0))
        
        # Segunda fila - Botones secundarios
        boton_ayuda = ttk.Button(
            frame_botones_bottom,
            text="❓ Ayuda & Tips",
            command=self.mostrar_ayuda,
            style='Action.TButton'
        )
        boton_ayuda.pack(side=tk.LEFT, padx=(0, 15))
        
        # Estadísticas en tiempo real
        self.label_stats = ttk.Label(
            frame_botones_bottom,
            text="💬 0 mensajes en esta sesión",
            font=('Segoe UI', 9),
            foreground=self.colores['texto_secundario']
        )
        self.label_stats.pack(side=tk.LEFT, padx=(15, 0))
        
        # Indicador de estado IA mejorado
        self.label_ia = ttk.Label(
            frame_botones_bottom,
            text="✨ IA Gemini 2.0 Flash" if self.chatbot.usar_ia else "⚡ Modo Offline",
            foreground=self.colores['success'] if self.chatbot.usar_ia else self.colores['warning'],
            font=('Segoe UI', 9, 'bold')
        )
        self.label_ia.pack(side=tk.RIGHT)
    
    def crear_status_bar(self, parent):
        """Crea la barra de estado mejorada"""
        frame_status = ttk.Frame(parent)
        frame_status.grid(row=5, column=0, sticky=(tk.W, tk.E), pady=(10, 0))
        frame_status.columnconfigure(1, weight=1)
        
        # Icono de estado
        self.status_icono = ttk.Label(
            frame_status,
            text="🟢",
            font=('Segoe UI Emoji', 12)
        )
        self.status_icono.grid(row=0, column=0, padx=(0, 8))
        
        # Mensaje de estado principal
        self.status_var = tk.StringVar()
        self.status_var.set("🚀 Sistema listo - ¡Comienza a conversar!")
        
        status_label = ttk.Label(
            frame_status,
            textvariable=self.status_var,
            font=('Segoe UI', 10, 'bold'),
            foreground=self.colores['texto_principal']
        )
        status_label.grid(row=0, column=1, sticky=(tk.W))
        
        # Información del sistema
        info_sistema = f"Python AI Assistant v2.0 | {datetime.now().strftime('%d/%m/%Y')}"
        info_label = ttk.Label(
            frame_status,
            text=info_sistema,
            font=('Segoe UI', 8),
            foreground=self.colores['texto_secundario']
        )
        info_label.grid(row=0, column=2, sticky=(tk.E))
    
    def cambiar_tema_animado(self):
        """Cambia entre modo claro y oscuro con animación"""
        # Cambiar estado
        self.modo_oscuro = not self.modo_oscuro
        
        # Actualizar texto del botón con animación
        nuevo_texto = "☀️ Modo Claro" if self.modo_oscuro else "🌙 Modo Oscuro"
        self.boton_tema.configure(text="🔄 Cambiando...")
        
        # Simular animación con after
        def aplicar_cambios():
            self.boton_tema.configure(text=nuevo_texto)
            self.aplicar_tema()
            self.configurar_estilos_dinamicos()
            self.actualizar_todos_los_componentes()
            
            # Mensaje de confirmación animado
            tema_nombre = "Oscuro 🌙" if self.modo_oscuro else "Claro ☀️"
            self.status_var.set(f"✨ Tema {tema_nombre} aplicado exitosamente")
            
            # Regresar al estado normal después de 2 segundos
            self.ventana.after(2000, lambda: self.status_var.set("Listo para conversar"))
        
        # Aplicar cambios después de un breve delay para efecto visual
        self.ventana.after(300, aplicar_cambios)
    
    def actualizar_todos_los_componentes(self):
        """Actualiza todos los componentes con el nuevo tema"""
        # Actualizar área de chat
        self.area_chat.configure(
            bg=self.colores['fondo_chat'],
            fg=self.colores['texto_principal'],
            insertbackground=self.colores['texto_principal'],
            selectbackground=self.colores['borde']
        )
        
        # Actualizar área de entrada de texto
        if hasattr(self, 'entrada_texto'):
            self.entrada_texto.configure(
                bg=self.colores['fondo_entry'],
                fg=self.colores['texto_principal'],
                insertbackground=self.colores['texto_principal'],
                selectbackground=self.colores['borde']
            )
        
        # Reconfigurar tags del chat
        self.configurar_tags_chat()
        
        # Actualizar colores del área de adjuntos
        self.actualizar_colores_adjuntos()
        
        # Actualizar indicadores de estado
        if hasattr(self, 'label_ia'):
            self.label_ia.configure(
                foreground=self.colores['success'] if self.chatbot.usar_ia else self.colores['warning']
            )
        
        if hasattr(self, 'indicador_ia'):
            self.indicador_ia.configure(
                foreground=self.colores['success'] if self.chatbot.usar_ia else self.colores['warning']
            )
        
        if hasattr(self, 'label_stats'):
            self.label_stats.configure(foreground=self.colores['texto_secundario'])
    
    def actualizar_estadisticas(self):
        """Actualiza las estadísticas en tiempo real"""
        if hasattr(self, 'label_stats'):
            total_mensajes = len(self.chatbot.sesion_actual['conversaciones'])
            if total_mensajes == 0:
                texto = "💬 Listo para conversar"
            elif total_mensajes == 1:
                texto = "💬 1 mensaje en esta sesión"
            else:
                texto = f"💬 {total_mensajes} mensajes en esta sesión"
            
            self.label_stats.configure(text=texto)
    
    def cambiar_tema(self):
        """Método de compatibilidad - redirige al método animado"""
        self.cambiar_tema_animado()
    
    def adjuntar_archivo(self):
        """Permite adjuntar archivos o imágenes"""
        tipos_archivo = [
            ("Todos los archivos soportados", "*.txt;*.pdf;*.docx;*.doc;*.png;*.jpg;*.jpeg;*.gif;*.bmp;*.json;*.xml;*.csv"),
            ("Documentos de texto", "*.txt;*.pdf;*.docx;*.doc"),
            ("Imágenes", "*.png;*.jpg;*.jpeg;*.gif;*.bmp"),
            ("Datos", "*.json;*.xml;*.csv"),
            ("Todos los archivos", "*.*")
        ]
        
        archivos = filedialog.askopenfilenames(
            title="Seleccionar archivos para adjuntar",
            filetypes=tipos_archivo
        )
        
        if archivos:
            for archivo in archivos:
                if archivo not in self.archivos_adjuntos:
                    self.archivos_adjuntos.append(archivo)
                    nombre_archivo = os.path.basename(archivo)
                    
                    # Determinar el tipo de archivo
                    extension = os.path.splitext(archivo)[1].lower()
                    if extension in ['.png', '.jpg', '.jpeg', '.gif', '.bmp']:
                        icono = "🖼️"
                    elif extension in ['.pdf']:
                        icono = "📄"
                    elif extension in ['.docx', '.doc']:
                        icono = "📝"
                    elif extension in ['.txt']:
                        icono = "📄"
                    elif extension in ['.json', '.xml']:
                        icono = "📋"
                    elif extension in ['.csv']:
                        icono = "📊"
                    else:
                        icono = "📎"
                    
                    self.lista_adjuntos.insert(tk.END, f"{icono} {nombre_archivo}")
            
            # Mostrar el área de adjuntos si hay archivos
            if self.archivos_adjuntos:
                self.frame_adjuntos.grid()
                self.actualizar_colores_adjuntos()
            
            self.status_var.set(f"📎 {len(self.archivos_adjuntos)} archivo(s) adjunto(s)")
    
    def limpiar_adjuntos_y_texto(self):
        """Limpia todos los archivos adjuntos sin tocar el texto del usuario"""
        # Limpiar archivos temporales de imágenes pegadas
        self.limpiar_archivos_temporales_actuales()
        
        # Limpiar lista de adjuntos
        self.archivos_adjuntos.clear()
        self.lista_adjuntos.delete(0, tk.END)
        
        # Ocultar frame de adjuntos
        self.frame_adjuntos.grid_remove()
        self.status_var.set("Listo para conversar")
    
    def limpiar_archivos_temporales_actuales(self):
        """Limpia solo los archivos temporales de la sesión actual"""
        try:
            for archivo in self.archivos_adjuntos[:]:  # Copia la lista para iterar
                if "imagen_pegada_" in os.path.basename(archivo):
                    try:
                        if os.path.exists(archivo):
                            os.remove(archivo)
                    except:
                        pass  # Si no se puede eliminar, continuar
        except:
            pass
    
    def quitar_archivo_adjunto(self):
        """Quita el archivo seleccionado de la lista"""
        seleccion = self.lista_adjuntos.curselection()
        if seleccion:
            indice = seleccion[0]
            archivo_a_quitar = self.archivos_adjuntos[indice]
            
            # Si es un archivo temporal de imagen pegada, eliminarlo del disco
            if "imagen_pegada_" in os.path.basename(archivo_a_quitar):
                try:
                    if os.path.exists(archivo_a_quitar):
                        os.remove(archivo_a_quitar)
                except:
                    pass
            
            # Quitar de la lista
            self.lista_adjuntos.delete(indice)
            del self.archivos_adjuntos[indice]
            
            if not self.archivos_adjuntos:
                self.frame_adjuntos.grid_remove()
                self.status_var.set("Listo para conversar")
            else:
                self.status_var.set(f"📎 {len(self.archivos_adjuntos)} archivo(s) adjunto(s)")
    
    def limpiar_adjuntos(self):
        """Limpia todos los archivos adjuntos"""
        self.archivos_adjuntos.clear()
        self.lista_adjuntos.delete(0, tk.END)
        self.frame_adjuntos.grid_remove()
        self.status_var.set("Listo para conversar")
    
    def actualizar_colores_adjuntos(self):
        """Actualiza los colores del área de adjuntos según el tema"""
        if hasattr(self, 'lista_adjuntos'):
            self.lista_adjuntos.configure(
                bg=self.colores['fondo_chat'],
                fg=self.colores['texto_principal'],
                selectbackground=self.colores['borde']
            )
    
    def procesar_archivos_adjuntos(self):
        """Procesa los archivos adjuntos y extrae su contenido"""
        contenido_archivos = []
        
        for archivo in self.archivos_adjuntos:
            try:
                nombre = os.path.basename(archivo)
                extension = os.path.splitext(archivo)[1].lower()
                
                if extension == '.txt':
                    with open(archivo, 'r', encoding='utf-8') as f:
                        contenido = f.read()
                    contenido_archivos.append(f"📄 Archivo: {nombre}\nContenido:\n{contenido}\n{'='*50}\n")
                
                elif extension == '.docx':
                    if DOCX_DISPONIBLE:
                        try:
                            doc = Document(archivo)
                            contenido = ""
                            for paragraph in doc.paragraphs:
                                if paragraph.text.strip():
                                    contenido += paragraph.text + "\n"
                            
                            # También extraer texto de tablas si las hay
                            for table in doc.tables:
                                for row in table.rows:
                                    for cell in row.cells:
                                        if cell.text.strip():
                                            contenido += cell.text + " | "
                                    contenido += "\n"
                            
                            if contenido.strip():
                                contenido_archivos.append(f"📝 Documento Word: {nombre}\nContenido:\n{contenido}\n{'='*50}\n")
                            else:
                                contenido_archivos.append(f"📝 Documento Word: {nombre}\n[El documento parece estar vacío o no contiene texto extraíble]\n{'='*50}\n")
                        except Exception as e:
                            contenido_archivos.append(f"❌ Error al leer documento Word {nombre}: {str(e)}\n{'='*50}\n")
                    else:
                        contenido_archivos.append(f"📝 Documento Word: {nombre}\n[No se puede leer - librería python-docx no disponible]\n{'='*50}\n")
                
                elif extension == '.doc':
                    # Para archivos .doc (formato antiguo), sugerir conversión
                    contenido_archivos.append(f"📝 Documento Word (formato antiguo): {nombre}\n[Por favor, convierte el archivo a formato .docx para poder leerlo, o copia y pega el contenido directamente]\n{'='*50}\n")
                
                elif extension == '.pdf':
                    if PDF_DISPONIBLE:
                        try:
                            contenido = self.extraer_texto_pdf(archivo)
                            if contenido.strip():
                                contenido_archivos.append(f"📄 Archivo PDF: {nombre}\nContenido:\n{contenido}\n{'='*50}\n")
                            else:
                                contenido_archivos.append(f"📄 Archivo PDF: {nombre}\n[No se pudo extraer texto del PDF o está vacío]\n{'='*50}\n")
                        except Exception as e:
                            contenido_archivos.append(f"❌ Error al leer PDF {nombre}: {str(e)}\n{'='*50}\n")
                    else:
                        contenido_archivos.append(f"📄 Archivo PDF: {nombre}\n[No se puede leer - librería PyPDF2 no disponible]\n{'='*50}\n")
                
                elif extension in ['.png', '.jpg', '.jpeg', '.gif', '.bmp']:
                    # Para imágenes, solo indicamos que están adjuntas
                    contenido_archivos.append(f"🖼️ Imagen adjunta: {nombre}\n[La imagen ha sido adjuntada para análisis]\n{'='*50}\n")
                
                elif extension == '.json':
                    with open(archivo, 'r', encoding='utf-8') as f:
                        contenido = f.read()
                    contenido_archivos.append(f"📋 Archivo JSON: {nombre}\nContenido:\n{contenido}\n{'='*50}\n")
                
                elif extension == '.csv':
                    with open(archivo, 'r', encoding='utf-8') as f:
                        contenido = f.read()
                    contenido_archivos.append(f"📊 Archivo CSV: {nombre}\nContenido:\n{contenido}\n{'='*50}\n")
                
                else:
                    # Para otros tipos de archivo, intentar leer como texto
                    try:
                        with open(archivo, 'r', encoding='utf-8') as f:
                            contenido = f.read()
                        contenido_archivos.append(f"📎 Archivo: {nombre}\nContenido:\n{contenido}\n{'='*50}\n")
                    except:
                        contenido_archivos.append(f"📎 Archivo: {nombre}\n[Archivo binario - no se puede mostrar el contenido]\n{'='*50}\n")
                        
            except Exception as e:
                contenido_archivos.append(f"❌ Error al leer {nombre}: {str(e)}\n{'='*50}\n")
        
        return "\n".join(contenido_archivos)
    
    def extraer_texto_pdf(self, ruta_pdf):
        """Extrae texto de un archivo PDF"""
        if not PDF_DISPONIBLE or PyPDF2 is None:
            return "Error: Librería PyPDF2 no disponible"
        
        try:
            texto = ""
            with open(ruta_pdf, 'rb') as archivo:
                # Usar la sintaxis correcta según la versión de PyPDF2
                try:
                    # Intentar con PyPDF2 versión 3.x
                    lector = PyPDF2.PdfReader(archivo)
                    for pagina in lector.pages:
                        texto += pagina.extract_text() + "\n"
                except AttributeError:
                    # Fallback para versiones anteriores de PyPDF2
                    lector = PyPDF2.PdfFileReader(archivo)
                    for num_pagina in range(lector.numPages):
                        pagina = lector.getPage(num_pagina)
                        texto += pagina.extractText() + "\n"
            return texto
        except Exception as e:
            return f"Error al extraer texto del PDF: {str(e)}"
    
    def mostrar_mensaje(self, mensaje, tipo, timestamp=True):
        """Muestra un mensaje en el área de chat"""
        self.area_chat.configure(state=tk.NORMAL)
        
        if timestamp:
            tiempo = datetime.now().strftime("%H:%M:%S")
            self.area_chat.insert(tk.END, f"[{tiempo}] ", "timestamp")
        
        if tipo == "usuario":
            self.area_chat.insert(tk.END, f"Tú: {mensaje}\n", "usuario")
        else:
            self.area_chat.insert(tk.END, f"🤖 Bot: {mensaje}\n", "bot")
        
        self.area_chat.insert(tk.END, "\n")
        self.area_chat.configure(state=tk.DISABLED)
        self.area_chat.see(tk.END)
    
    def mostrar_mensaje_bot(self, mensaje):
        """Muestra un mensaje del bot"""
        self.mostrar_mensaje(mensaje, "bot")
    
    def mostrar_mensaje_usuario(self, mensaje):
        """Muestra un mensaje del usuario"""
        self.mostrar_mensaje(mensaje, "usuario")
    
    def configurar_menu_contextual(self):
        """Configura el menú contextual (clic derecho) para el área de entrada"""
        self.menu_contextual = tk.Menu(self.entrada_texto, tearoff=0)
        
        # Opciones del menú
        self.menu_contextual.add_command(label="🔄 Deshacer", command=self.deshacer, accelerator="Ctrl+Z")
        self.menu_contextual.add_command(label="↩️ Rehacer", command=self.rehacer, accelerator="Ctrl+Y")
        self.menu_contextual.add_separator()
        self.menu_contextual.add_command(label="✂️ Cortar", command=self.cortar, accelerator="Ctrl+X")
        self.menu_contextual.add_command(label="📋 Copiar", command=self.copiar, accelerator="Ctrl+C")
        self.menu_contextual.add_command(label="📝 Pegar texto/imagen", command=self.pegar, accelerator="Ctrl+V")
        self.menu_contextual.add_separator()
        self.menu_contextual.add_command(label="🎯 Seleccionar todo", command=self.seleccionar_todo, accelerator="Ctrl+A")
        self.menu_contextual.add_command(label="🗑️ Limpiar texto", command=self.limpiar_entrada, accelerator="Ctrl+L")
        self.menu_contextual.add_command(label="🧹 Limpiar adjuntos", command=self.limpiar_adjuntos_y_texto, accelerator="Ctrl+Shift+L")
        
        # Bind clic derecho
        self.entrada_texto.bind("<Button-3>", self.mostrar_menu_contextual)
    
    def mostrar_menu_contextual(self, event):
        """Muestra el menú contextual en la posición del cursor"""
        try:
            # Actualizar estado de opciones según contexto
            tiene_seleccion = False
            try:
                if self.entrada_texto.selection_get():
                    tiene_seleccion = True
            except tk.TclError:
                pass
            
            tiene_texto = len(self.entrada_texto.get("1.0", tk.END).strip()) > 0
            
            # Habilitar/deshabilitar opciones según contexto
            self.menu_contextual.entryconfig("✂️ Cortar", state="normal" if tiene_seleccion else "disabled")
            self.menu_contextual.entryconfig("📋 Copiar", state="normal" if tiene_seleccion else "disabled")
            self.menu_contextual.entryconfig("🗑️ Limpiar", state="normal" if tiene_texto else "disabled")
            
            # Verificar si hay algo en el portapapeles (texto o imagen)
            puede_pegar = False
            try:
                # Verificar texto
                self.entrada_texto.clipboard_get()
                puede_pegar = True
            except tk.TclError:
                pass
            
            # Verificar imagen en portapapeles si PIL está disponible
            if not puede_pegar and PIL_DISPONIBLE:
                try:
                    imagen = ImageGrab.grabclipboard()
                    if imagen is not None and isinstance(imagen, Image.Image):
                        puede_pegar = True
                except:
                    pass
            
            self.menu_contextual.entryconfig("📝 Pegar", state="normal" if puede_pegar else "disabled")
            
            # Mostrar menú
            self.menu_contextual.post(event.x_root, event.y_root)
        except Exception as e:
            print(f"Error en menú contextual: {e}")
    
    def configurar_atajos_teclado(self):
        """Configura los atajos de teclado estándar para el área de entrada"""
        # Atajos de edición estándar
        self.entrada_texto.bind('<Control-z>', self.deshacer)
        self.entrada_texto.bind('<Control-Z>', self.deshacer)
        self.entrada_texto.bind('<Control-y>', self.rehacer)
        self.entrada_texto.bind('<Control-Y>', self.rehacer)
        
        # Atajos de portapapeles
        self.entrada_texto.bind('<Control-c>', self.copiar)
        self.entrada_texto.bind('<Control-C>', self.copiar)
        self.entrada_texto.bind('<Control-v>', self.pegar)
        self.entrada_texto.bind('<Control-V>', self.pegar)
        self.entrada_texto.bind('<Control-x>', self.cortar)
        self.entrada_texto.bind('<Control-X>', self.cortar)
        
        # Atajos de selección
        self.entrada_texto.bind('<Control-a>', self.seleccionar_todo)
        self.entrada_texto.bind('<Control-A>', self.seleccionar_todo)
        
        # Atajos adicionales útiles
        self.entrada_texto.bind('<Control-l>', self.limpiar_entrada)
        self.entrada_texto.bind('<Control-L>', self.limpiar_entrada)
        self.entrada_texto.bind('<Control-Shift-L>', self.limpiar_adjuntos_y_texto)
        self.entrada_texto.bind('<Control-Shift-l>', self.limpiar_adjuntos_y_texto)
        
        # Navegación rápida
        self.entrada_texto.bind('<Control-Home>', self.ir_inicio)
        self.entrada_texto.bind('<Control-End>', self.ir_final)
    
    def deshacer(self, event=None):
        """Deshace la última acción"""
        try:
            self.entrada_texto.edit_undo()
        except tk.TclError:
            pass  # No hay nada que deshacer
        return "break"
    
    def rehacer(self, event=None):
        """Rehace la última acción deshecha"""
        try:
            self.entrada_texto.edit_redo()
        except tk.TclError:
            pass  # No hay nada que rehacer
        return "break"
    
    def copiar(self, event=None):
        """Copia el texto seleccionado al portapapeles"""
        try:
            if self.entrada_texto.selection_get():
                self.entrada_texto.clipboard_clear()
                self.entrada_texto.clipboard_append(self.entrada_texto.selection_get())
        except tk.TclError:
            pass  # No hay selección
        return "break"
    
    def pegar(self, event=None):
        """Pega texto o imágenes del portapapeles"""
        try:
            # Primero intentar pegar imagen del portapapeles
            if PIL_DISPONIBLE and self.pegar_imagen_portapapeles():
                return "break"
            
            # Si no hay imagen, pegar texto normal
            cursor_pos = self.entrada_texto.index(tk.INSERT)
            texto_portapapeles = self.entrada_texto.clipboard_get()
            self.entrada_texto.insert(cursor_pos, texto_portapapeles)
        except tk.TclError:
            pass  # Portapapeles vacío o error
        return "break"
    
    def pegar_imagen_portapapeles(self):
        """Intenta pegar una imagen del portapapeles como archivo adjunto"""
        try:
            if not PIL_DISPONIBLE:
                return False
            
            # Intentar obtener imagen del portapapeles
            imagen = ImageGrab.grabclipboard()
            
            if imagen is not None and isinstance(imagen, Image.Image):
                # Crear nombre único para el archivo temporal
                timestamp = datetime.now().strftime("%Y%m%d_%H%M%S_%f")[:-3]  # Incluir microsegundos
                nombre_archivo = f"imagen_pegada_{timestamp}.png"
                
                # Crear directorio temporal si no existe
                temp_dir = os.path.join(tempfile.gettempdir(), "chatbot_imagenes")
                os.makedirs(temp_dir, exist_ok=True)
                
                # Guardar imagen como archivo temporal
                ruta_temporal = os.path.join(temp_dir, nombre_archivo)
                imagen.save(ruta_temporal, "PNG")
                
                # Verificar que el archivo no esté ya en la lista
                if ruta_temporal not in self.archivos_adjuntos:
                    self.archivos_adjuntos.append(ruta_temporal)
                    self.lista_adjuntos.insert(tk.END, f"🖼️ {nombre_archivo} (pegada)")
                    
                    # Mostrar el área de adjuntos
                    self.frame_adjuntos.grid()
                    self.actualizar_colores_adjuntos()
                    
                    # Actualizar status
                    self.status_var.set(f"📎 Imagen pegada - {len(self.archivos_adjuntos)} archivo(s) adjunto(s)")
                    
                    # NO insertar texto de referencia en el área de texto
                
                return True
            
            return False
            
        except Exception as e:
            print(f"Error al pegar imagen: {e}")
            return False
    
    def cortar(self, event=None):
        """Corta el texto seleccionado"""
        try:
            if self.entrada_texto.selection_get():
                # Copiar al portapapeles
                self.entrada_texto.clipboard_clear()
                self.entrada_texto.clipboard_append(self.entrada_texto.selection_get())
                # Eliminar selección
                self.entrada_texto.delete(tk.SEL_FIRST, tk.SEL_LAST)
        except tk.TclError:
            pass  # No hay selección
        return "break"
    
    def seleccionar_todo(self, event=None):
        """Selecciona todo el texto"""
        self.entrada_texto.tag_add(tk.SEL, "1.0", tk.END)
        self.entrada_texto.mark_set(tk.INSERT, "1.0")
        self.entrada_texto.see(tk.INSERT)
        return "break"
    
    def limpiar_entrada(self, event=None):
        """Limpia toda el área de entrada (Ctrl+L)"""
        self.entrada_texto.delete("1.0", tk.END)
        return "break"
    
    def ir_inicio(self, event=None):
        """Va al inicio del texto (Ctrl+Home)"""
        self.entrada_texto.mark_set(tk.INSERT, "1.0")
        self.entrada_texto.see(tk.INSERT)
        return "break"
    
    def ir_final(self, event=None):
        """Va al final del texto (Ctrl+End)"""
        self.entrada_texto.mark_set(tk.INSERT, tk.END)
        self.entrada_texto.see(tk.INSERT)
        return "break"
    
    def manejar_enter(self, event):
        """Maneja la tecla Enter (envía el mensaje)"""
        self.enviar_mensaje()
        return "break"  # Evita que se agregue una nueva línea
    
    def manejar_shift_enter(self, event):
        """Maneja Shift+Enter (nueva línea)"""
        # No hacer nada especial, permitir el comportamiento por defecto
        return None
    
    def enviar_mensaje(self):
        """Envía un mensaje al chatbot"""
        # Obtener todo el contenido del Text widget
        mensaje = self.entrada_texto.get("1.0", tk.END).strip()
        
        # Verificar si hay mensaje o archivos adjuntos
        if not mensaje and not self.archivos_adjuntos:
            return
        
        # Si no hay mensaje pero sí archivos, crear mensaje automático
        if not mensaje and self.archivos_adjuntos:
            mensaje = "Por favor, analiza los archivos adjuntos."
        
        # Limpiar entrada
        self.entrada_texto.delete("1.0", tk.END)
        
        # Procesar archivos adjuntos si los hay
        contenido_adjuntos = ""
        if self.archivos_adjuntos:
            contenido_adjuntos = self.procesar_archivos_adjuntos()
            # Mostrar información de archivos adjuntos en el chat
            self.mostrar_mensaje_adjuntos()
        
        # Combinar mensaje con contenido de archivos
        mensaje_completo = mensaje
        if contenido_adjuntos:
            mensaje_completo += f"\n\n--- ARCHIVOS ADJUNTOS ---\n{contenido_adjuntos}"
        
        # Mostrar mensaje del usuario (solo el texto, no los archivos)
        self.mostrar_mensaje_usuario(mensaje)
        
        # Verificar si quiere salir
        if mensaje.lower() in ['salir', 'exit', 'quit']:
            self.cerrar_aplicacion()
            return
        
        # Deshabilitar botón mientras procesa
        self.boton_enviar.configure(state='disabled')
        self.boton_adjuntar.configure(state='disabled')
        
        # Status dinámico según el tipo de consulta
        if self.archivos_adjuntos:
            self.status_var.set("🔍 Analizando archivos adjuntos con IA...")
            self.status_icono.configure(text="🔄")
        else:
            self.status_var.set("🤔 Generando respuesta inteligente...")
            self.status_icono.configure(text="🧠")
        
        # Procesar mensaje en hilo separado
        threading.Thread(
            target=self.procesar_mensaje_async,
            args=(mensaje_completo,),
            daemon=True
        ).start()
        
        # Limpiar archivos adjuntos después de enviar
        if self.archivos_adjuntos:
            self.limpiar_adjuntos()
    
    def mostrar_mensaje_adjuntos(self):
        """Muestra información sobre los archivos adjuntos en el chat"""
        if self.archivos_adjuntos:
            archivos_info = []
            for archivo in self.archivos_adjuntos:
                nombre = os.path.basename(archivo)
                extension = os.path.splitext(archivo)[1].lower()
                
                if extension in ['.png', '.jpg', '.jpeg', '.gif', '.bmp']:
                    icono = "🖼️"
                elif extension == '.pdf':
                    icono = "📄"
                elif extension in ['.docx', '.doc']:
                    icono = "📝"
                elif extension == '.txt':
                    icono = "📄"
                elif extension in ['.json', '.xml']:
                    icono = "📋"
                elif extension == '.csv':
                    icono = "📊"
                else:
                    icono = "📎"
                
                archivos_info.append(f"{icono} {nombre}")
            
            mensaje_adjuntos = f"📎 Archivos adjuntos: {', '.join(archivos_info)}"
            self.mostrar_mensaje(mensaje_adjuntos, "usuario", timestamp=False)
    
    def procesar_mensaje_async(self, mensaje):
        """Procesa el mensaje en un hilo separado"""
        try:
            respuesta = self.chatbot.procesar_mensaje(mensaje)
            
            # Actualizar UI en el hilo principal
            self.ventana.after(0, self.mostrar_respuesta, respuesta)
            
        except Exception as e:
            error_msg = f"Lo siento, ocurrió un error: {str(e)}"
            self.ventana.after(0, self.mostrar_respuesta, error_msg)
    
    def mostrar_respuesta(self, respuesta):
        """Muestra la respuesta del bot y reactiva controles"""
        self.mostrar_mensaje_bot(respuesta)
        self.boton_enviar.configure(state='normal')
        self.boton_adjuntar.configure(state='normal')
        
        # Actualizar estadísticas
        self.actualizar_estadisticas()
        
        # Status con animación
        self.status_var.set("✨ Respuesta generada exitosamente")
        self.status_icono.configure(text="✅")
        
        # Regresar al estado normal después de 2 segundos
        self.ventana.after(2000, self.restablecer_status_normal)
        
        self.entrada_texto.focus()
    
    def restablecer_status_normal(self):
        """Restablece el status a normal"""
        self.status_var.set("🚀 Listo para la siguiente consulta")
        self.status_icono.configure(text="🟢")
    
    def limpiar_chat(self):
        """Limpia el área de chat"""
        respuesta = messagebox.askyesno(
            "Limpiar Chat", 
            "¿Estás seguro de que quieres limpiar toda la conversación?"
        )
        if respuesta:
            self.area_chat.configure(state=tk.NORMAL)
            self.area_chat.delete(1.0, tk.END)
            self.area_chat.configure(state=tk.DISABLED)
            self.mostrar_mensaje_bot("Chat limpiado. ¡Empecemos de nuevo! 😊")
    
    def abrir_ventana_historial(self):
        """Abre la ventana del historial de conversaciones"""
        ventana_historial = tk.Toplevel(self.ventana)
        ventana_historial.title("📂 Historial de Conversaciones")
        ventana_historial.geometry("900x700")
        ventana_historial.configure(bg=self.colores['fondo_principal'])
        
        # Frame principal
        frame_principal = ttk.Frame(ventana_historial, padding="10")
        frame_principal.pack(fill=tk.BOTH, expand=True)
        
        # Título
        titulo = ttk.Label(
            frame_principal, 
            text="📂 Historial de Conversaciones", 
            font=('Arial', 14, 'bold')
        )
        titulo.pack(pady=(0, 10))
        
        # Frame para lista y detalles
        frame_contenido = ttk.Frame(frame_principal)
        frame_contenido.pack(fill=tk.BOTH, expand=True)
        
        # Lista de sesiones (izquierda)
        frame_lista = ttk.LabelFrame(frame_contenido, text="📋 Sesiones", padding="5")
        frame_lista.pack(side=tk.LEFT, fill=tk.BOTH, expand=True, padx=(0, 5))
        
        # Treeview para sesiones
        columnas = ('Fecha', 'Hora', 'Mensajes')
        tree_sesiones = ttk.Treeview(frame_lista, columns=columnas, show='headings', height=15)
        
        # Configurar columnas
        tree_sesiones.heading('Fecha', text='Fecha')
        tree_sesiones.heading('Hora', text='Hora')
        tree_sesiones.heading('Mensajes', text='Mensajes')
        
        tree_sesiones.column('Fecha', width=100)
        tree_sesiones.column('Hora', width=80)
        tree_sesiones.column('Mensajes', width=80)
        
        # Scrollbar para treeview
        scroll_tree = ttk.Scrollbar(frame_lista, orient=tk.VERTICAL, command=tree_sesiones.yview)
        tree_sesiones.configure(yscrollcommand=scroll_tree.set)
        
        tree_sesiones.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        scroll_tree.pack(side=tk.RIGHT, fill=tk.Y)
        
        # Área de detalles (derecha)
        frame_detalles = ttk.LabelFrame(frame_contenido, text="💬 Conversación", padding="5")
        frame_detalles.pack(side=tk.RIGHT, fill=tk.BOTH, expand=True, padx=(5, 0))
        
        # Área de texto para mostrar conversación
        area_conversacion = scrolledtext.ScrolledText(
            frame_detalles,
            wrap=tk.WORD,
            state=tk.DISABLED,
            font=('Segoe UI', 10),
            bg=self.colores['fondo_chat'],
            fg=self.colores['texto_principal'],
            insertbackground=self.colores['texto_principal']
        )
        area_conversacion.pack(fill=tk.BOTH, expand=True)
        
        # Configurar tags para colores (adaptable a tema)
        area_conversacion.tag_configure("usuario", 
                                       foreground=self.colores['texto_usuario'], 
                                       font=('Segoe UI', 10, 'bold'),
                                       justify='left')
        area_conversacion.tag_configure("bot", 
                                       foreground=self.colores['texto_bot'], 
                                       font=('Segoe UI', 10),
                                       justify='left')
        area_conversacion.tag_configure("timestamp", 
                                       foreground=self.colores['texto_secundario'], 
                                       font=('Segoe UI', 8),
                                       justify='left')
        area_conversacion.tag_configure("ia_badge", 
                                       foreground=self.colores['texto_secundario'], 
                                       font=('Segoe UI', 8, 'italic'),
                                       justify='left')
        
        # Cargar sesiones
        self.cargar_sesiones_en_tree(tree_sesiones)
        
        # Evento de selección
        def on_select(event):
            selection = tree_sesiones.selection()
            if selection:
                item = tree_sesiones.item(selection[0])
                archivo = item['values'][3]  # Archivo guardado en valores ocultos
                self.mostrar_conversacion_en_area(archivo, area_conversacion)
        
        tree_sesiones.bind('<<TreeviewSelect>>', on_select)
        
        # Frame para botones
        frame_botones_hist = ttk.Frame(frame_principal)
        frame_botones_hist.pack(fill=tk.X, pady=(10, 0))
        
        # Botón actualizar
        btn_actualizar = ttk.Button(
            frame_botones_hist,
            text="🔄 Actualizar",
            command=lambda: self.cargar_sesiones_en_tree(tree_sesiones)
        )
        btn_actualizar.pack(side=tk.LEFT, padx=(0, 10))
        
        # Botón exportar seleccionada
        btn_exportar_sel = ttk.Button(
            frame_botones_hist,
            text="💾 Exportar Seleccionada",
            command=lambda: self.exportar_sesion_seleccionada(tree_sesiones)
        )
        btn_exportar_sel.pack(side=tk.LEFT, padx=(0, 10))
        
        # Estadísticas
        stats = self.chatbot.obtener_estadisticas_historial()
        if stats:
            texto_stats = f"📊 Total: {stats['total_sesiones']} sesiones, {stats['total_conversaciones']} mensajes"
            label_stats = ttk.Label(frame_botones_hist, text=texto_stats)
            label_stats.pack(side=tk.RIGHT)
    
    def cargar_sesiones_en_tree(self, tree):
        """Carga las sesiones en el treeview"""
        # Limpiar tree
        for item in tree.get_children():
            tree.delete(item)
        
        # Cargar sesiones
        sesiones = self.chatbot.cargar_historial_sesiones()
        
        for sesion in sesiones:
            datos = sesion['datos']
            inicio = datetime.fromisoformat(datos['inicio'])
            fecha = inicio.strftime("%Y-%m-%d")
            hora = inicio.strftime("%H:%M:%S")
            mensajes = len(datos['conversaciones'])
            
            # Insertar en tree (guardamos el archivo en values[3] oculto)
            tree.insert('', tk.END, values=(fecha, hora, mensajes, sesion['archivo']))
    
    def mostrar_conversacion_en_area(self, archivo, area):
        """Muestra una conversación específica en el área de texto"""
        try:
            ruta_archivo = os.path.join(self.chatbot.directorio_historial, archivo)
            
            with open(ruta_archivo, 'r', encoding='utf-8') as f:
                sesion = json.load(f)
            
            area.configure(state=tk.NORMAL)
            area.delete(1.0, tk.END)
            
            # Información de la sesión
            inicio = datetime.fromisoformat(sesion['inicio'])
            area.insert(tk.END, f"📅 Sesión del {inicio.strftime('%Y-%m-%d %H:%M:%S')}\n", "timestamp")
            area.insert(tk.END, f"💬 Total de mensajes: {len(sesion['conversaciones'])}\n\n", "timestamp")
            
            # Mostrar conversaciones
            for i, conv in enumerate(sesion['conversaciones'], 1):
                timestamp = datetime.fromisoformat(conv['timestamp'])
                tiempo = timestamp.strftime("%H:%M:%S")
                
                # Badge de IA si fue respuesta de IA
                ia_badge = " [IA]" if conv.get('fue_ia', False) else " [Local]"
                
                area.insert(tk.END, f"[{tiempo}] ", "timestamp")
                area.insert(tk.END, f"Tú: {conv['usuario']}\n", "usuario")
                
                area.insert(tk.END, f"[{tiempo}] ", "timestamp")
                area.insert(tk.END, f"🤖 Bot{ia_badge}: {conv['bot']}\n\n", "bot")
            
            area.configure(state=tk.DISABLED)
            area.see(1.0)  # Ir al inicio
            
        except Exception as e:
            messagebox.showerror("Error", f"No se pudo cargar la conversación: {e}")
    
    def exportar_sesion_seleccionada(self, tree):
        """Exporta la sesión seleccionada a un archivo de texto"""
        selection = tree.selection()
        if not selection:
            messagebox.showwarning("Advertencia", "Por favor selecciona una sesión para exportar.")
            return
        
        item = tree.item(selection[0])
        archivo = item['values'][3]
        
        try:
            ruta_archivo = os.path.join(self.chatbot.directorio_historial, archivo)
            
            with open(ruta_archivo, 'r', encoding='utf-8') as f:
                sesion = json.load(f)
            
            # Crear contenido de texto
            inicio = datetime.fromisoformat(sesion['inicio'])
            contenido = f"Conversación con Asistente Virtual\n"
            contenido += f"Fecha: {inicio.strftime('%Y-%m-%d %H:%M:%S')}\n"
            contenido += f"Total de mensajes: {len(sesion['conversaciones'])}\n"
            contenido += "=" * 50 + "\n\n"
            
            for conv in sesion['conversaciones']:
                timestamp = datetime.fromisoformat(conv['timestamp'])
                tiempo = timestamp.strftime("%H:%M:%S")
                ia_badge = " [IA]" if conv.get('fue_ia', False) else " [Local]"
                
                contenido += f"[{tiempo}] Usuario: {conv['usuario']}\n"
                contenido += f"[{tiempo}] Bot{ia_badge}: {conv['bot']}\n\n"
            
            # Guardar archivo
            archivo_export = filedialog.asksaveasfilename(
                title="Guardar conversación",
                defaultextension=".txt",
                filetypes=[("Archivos de texto", "*.txt"), ("Todos los archivos", "*.*")]
            )
            
            if archivo_export:
                with open(archivo_export, 'w', encoding='utf-8') as f:
                    f.write(contenido)
                messagebox.showinfo("Éxito", "Conversación exportada exitosamente.")
                
        except Exception as e:
            messagebox.showerror("Error", f"No se pudo exportar la conversación: {e}")
    
    def guardar_conversacion_actual(self):
        """Guarda la conversación actual en el historial"""
        if len(self.chatbot.sesion_actual['conversaciones']) == 0:
            messagebox.showwarning("Advertencia", "No hay conversación actual para guardar.")
            return
        
        try:
            # Guardar la sesión actual en el historial
            archivo_guardado = self.chatbot.guardar_sesion_completa()
            
            if archivo_guardado:
                messagebox.showinfo(
                    "Éxito", 
                    f"Conversación guardada exitosamente en el historial.\n\nArchivo: {archivo_guardado}"
                )
                self.status_var.set("✅ Conversación guardada en historial")
            else:
                messagebox.showwarning("Advertencia", "No se pudo guardar la conversación.")
                
        except Exception as e:
            messagebox.showerror("Error", f"No se pudo guardar la conversación: {e}")
    
    def exportar_conversacion_actual(self):
        """Exporta la conversación actual que se está viendo"""
        if len(self.chatbot.sesion_actual['conversaciones']) == 0:
            messagebox.showwarning("Advertencia", "No hay conversación actual para exportar.")
            return
        
        try:
            # Crear contenido de texto
            inicio = datetime.fromisoformat(self.chatbot.sesion_actual['inicio'])
            contenido = f"Conversación Actual con Asistente Virtual\n"
            contenido += f"Fecha: {inicio.strftime('%Y-%m-%d %H:%M:%S')}\n"
            contenido += f"Total de mensajes: {len(self.chatbot.sesion_actual['conversaciones'])}\n"
            contenido += "=" * 50 + "\n\n"
            
            for conv in self.chatbot.sesion_actual['conversaciones']:
                timestamp = datetime.fromisoformat(conv['timestamp'])
                tiempo = timestamp.strftime("%H:%M:%S")
                ia_badge = " [IA]" if conv.get('fue_ia', False) else " [Local]"
                
                contenido += f"[{tiempo}] Usuario: {conv['usuario']}\n"
                contenido += f"[{tiempo}] Bot{ia_badge}: {conv['bot']}\n\n"
            
            # Guardar archivo
            archivo_export = filedialog.asksaveasfilename(
                title="Guardar conversación actual",
                defaultextension=".txt",
                filetypes=[("Archivos de texto", "*.txt"), ("Todos los archivos", "*.*")]
            )
            
            if archivo_export:
                with open(archivo_export, 'w', encoding='utf-8') as f:
                    f.write(contenido)
                messagebox.showinfo("Éxito", "Conversación actual exportada exitosamente.")
                
        except Exception as e:
            messagebox.showerror("Error", f"No se pudo exportar la conversación: {e}")
    
    def mostrar_ayuda(self):
        """Muestra información de ayuda"""
        ayuda_texto = """
🤖 Asistente Virtual - Ayuda

ENTRADA DE TEXTO:
• Escribe tu pregunta y presiona Enter o clic en 'Enviar'
• Shift+Enter: Nueva línea para texto multilínea
• Escribe 'salir' para cerrar la aplicación

⌨️ ATAJOS DE TECLADO:
• Ctrl+Z: Deshacer última acción
• Ctrl+Y: Rehacer acción deshecha
• Ctrl+C: Copiar texto seleccionado
• Ctrl+V: Pegar texto o imágenes desde portapapeles
• Ctrl+X: Cortar texto seleccionado
• Ctrl+A: Seleccionar todo el texto
• Ctrl+L: Limpiar solo el área de entrada
• Ctrl+Shift+L: Limpiar solo archivos adjuntos
• Ctrl+Home: Ir al inicio del texto
• Ctrl+End: Ir al final del texto
• Clic derecho: Menú contextual con opciones

📋 PEGAR IMÁGENES:
• Copia cualquier imagen (Ctrl+C en navegador, captura de pantalla, etc.)
• Pega directamente en el área de texto con Ctrl+V
• La imagen se adjunta automáticamente sin alterar tu texto
• Usa "Limpiar Adjuntos" para eliminar solo las imágenes pegadas

Funciones disponibles:
✨ Respuestas inteligentes con IA
💬 Conversación natural
🧠 Memoria de contexto
🔄 Respuestas de respaldo si falla la IA
📂 Historial completo de conversaciones
💾 Guardar conversaciones en historial
🌙/☀️ Modo oscuro y claro
🎨 Interfaz elegante y adaptable

📎 ARCHIVOS ADJUNTOS:
• Adjunta documentos (.txt, .docx, .pdf)
• Adjunta imágenes (.png, .jpg, .jpeg, .gif, .bmp)
• Adjunta datos (.json, .xml, .csv)
• ✨ Lectura completa de archivos Word (.docx)
• Extracción automática de texto y tablas

🎯 SOLICITUDES ESPECÍFICAS:
• "Dame un resumen del documento" - Solo resumen
• "Genera casos de prueba" - Solo casos de prueba
• "Analiza el código" - Solo análisis de código
• "Revisa el documento" - Revisión general
• Sé específico en tu solicitud para mejores resultados

🎭 ROLES PROFESIONALES DISPONIBLES:
• "Actúa como experto en QA" - Especialista en testing y casos de prueba
• "Actúa como arquitecto de software" - Diseño de arquitecturas
• "Actúa como analista de negocio" - Análisis de procesos y requisitos
• "Actúa como desarrollador senior" - Revisión y mejores prácticas de código
• "Actúa como consultor técnico" - Guidance experto y recomendaciones

¡Pregúntame lo que quieras y adjunta archivos para análisis!
        """
        messagebox.showinfo("Ayuda - Asistente Virtual", ayuda_texto)

    def limpiar_archivos_temporales(self):
        """Limpia archivos temporales de imágenes pegadas"""
        try:
            temp_dir = os.path.join(tempfile.gettempdir(), "chatbot_imagenes")
            if os.path.exists(temp_dir):
                for archivo in os.listdir(temp_dir):
                    if archivo.startswith("imagen_pegada_"):
                        ruta_archivo = os.path.join(temp_dir, archivo)
                        try:
                            os.remove(ruta_archivo)
                        except:
                            pass  # Si no se puede eliminar, continuar
        except:
            pass  # Si hay error, no importa
    
    def cerrar_aplicacion(self):
        """Cierra la aplicación guardando la sesión"""
        try:
            # Limpiar archivos temporales
            self.limpiar_archivos_temporales()
            
            # Guardar sesión actual si hay conversaciones
            if len(self.chatbot.sesion_actual['conversaciones']) > 0:
                archivo_guardado = self.chatbot.guardar_sesion_completa()
                if archivo_guardado:
                    print(f"✅ Sesión guardada en: {archivo_guardado}")
            
            self.ventana.quit()
        except Exception as e:
            print(f"Error al guardar sesión: {e}")
            self.ventana.quit()
    
    def ejecutar(self):
        """Inicia la aplicación"""
        try:
            self.ventana.mainloop()
        except KeyboardInterrupt:
            self.cerrar_aplicacion()

def main():
    """Función principal"""
    app = ChatBotGUI()
    app.ejecutar()

if __name__ == "__main__":
    main()
